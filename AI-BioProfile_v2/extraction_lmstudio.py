import base64
import io
import json
import os
import re
import urllib.request
import zipfile
from urllib.parse import urlparse, urlunparse
from typing import Any, Dict, List, Optional

import fitz
import pytesseract
import requests
from requests.exceptions import ConnectionError, Timeout
from docx import Document
from PIL import Image

from test_extraction_v2 import CVExtraction


TESSERACT_DEFAULT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(TESSERACT_DEFAULT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_DEFAULT_PATH

LMSTUDIO_API_URL = os.getenv(
    "LMSTUDIO_API_URL",
    "http://127.0.0.1:1234/v1/chat/completions"
)
LMSTUDIO_MODEL = os.getenv("LMSTUDIO_MODEL", "local-model")
LMSTUDIO_CONNECT_TIMEOUT = float(os.getenv("LMSTUDIO_CONNECT_TIMEOUT", "10"))
LMSTUDIO_READ_TIMEOUT = float(os.getenv("LMSTUDIO_READ_TIMEOUT", "1800"))
LMSTUDIO_MAX_TOKENS = int(os.getenv("LMSTUDIO_MAX_TOKENS", "4096"))


def _to_base64_png(img: Image.Image) -> str:
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def pdf_pages_to_base64_images(file_bytes: bytes, max_pages: int = 2, zoom: float = 1.0) -> List[str]:
    """Render first PDF pages as PNG base64 for VLM input."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images: List[str] = []
    page_count = min(len(doc), max_pages)

    for page_index in range(page_count):
        page = doc[page_index]
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(_to_base64_png(img))

    return images


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Text fallback used alongside VLM images to improve extraction reliability."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text().strip()
        
        # Si le texte extrait est vide ou très court (potentiellement un PDF scanné)
        if len(page_text) < 20:
            try:
                # Zoom x2 pour améliorer la qualité de l'OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                # On tente l'OCR en français et en anglais
                ocr_text = pytesseract.image_to_string(img, lang='fra+eng').strip()
                text += ocr_text + "\n"
            except Exception as e:
                print(f"Erreur OCR sur la page: {e}")
                text += page_text + "\n"
        else:
            text += page_text + "\n"
            
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception:
        return ""


def extract_image_from_docx(file_bytes: bytes) -> Optional[bytes]:
    """Extract best probable profile image from DOCX media."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            media_files = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
            if not media_files:
                return None

            candidates = []
            for media_path in media_files:
                with archive.open(media_path) as media_file:
                    img_data = media_file.read()
                try:
                    image = Image.open(io.BytesIO(img_data))
                    width, height = image.size
                    if width < 50 or height < 50:
                        continue
                    ratio = width / height if height else 1.0
                    score = 100.0 - 30.0 * abs(1.0 - ratio)
                    if width > 1000 or height > 1000:
                        score -= 30.0
                    candidates.append((score, img_data))
                except Exception:
                    continue

            if not candidates:
                return None

            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
    except Exception:
        return None


def extract_image_from_pdf(file_bytes: bytes) -> Optional[bytes]:
    """Reuse deterministic image heuristics to preserve profile photo behavior."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if len(doc) == 0:
            return None

        candidates = []
        is_scanned_doc = False

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_rect = page.rect
            page_w = page_rect.width
            page_h = page_rect.height
            image_list = page.get_images(full=True)

            for img in image_list:
                xref = img[0]
                try:
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n >= 5:
                        pix = fitz.Pixmap(fitz.csRGB, pix)

                    raw_bytes = pix.tobytes("png")
                    img_pil = Image.open(io.BytesIO(raw_bytes))

                    if img_pil.mode in ("RGBA", "LA") or (img_pil.mode == "P" and "transparency" in img_pil.info):
                        alpha = img_pil.convert("RGBA").split()[-1]
                        bg = Image.new("RGB", img_pil.size, (255, 255, 255))
                        bg.paste(img_pil, mask=alpha)
                        img_pil = bg
                    elif img_pil.mode != "RGB":
                        img_pil = img_pil.convert("RGB")

                    output = io.BytesIO()
                    img_pil.save(output, format="PNG")
                    image_bytes = output.getvalue()

                    width = pix.width
                    height = pix.height
                    pix = None
                except Exception:
                    continue

                aspect_ratio = width / height if height != 0 else 1.0
                if aspect_ratio < 0.4 or aspect_ratio > 2.5:
                    continue

                rects = page.get_image_rects(xref)
                if rects:
                    r = rects[0]
                    if r.width > page_w * 0.75 or r.height > page_h * 0.75:
                        is_scanned_doc = True
                        continue
                    if r.width > page_w * 0.5 and r.height > page_h * 0.5:
                        is_scanned_doc = True
                        continue
                    if r.width < 25 or r.height < 25:
                        continue
                    if r.y0 > page_h * 0.7:
                        continue
                else:
                    if width < 100 or height < 100:
                        continue
                    if width > 900 or height > 900:
                        continue

                score = 100.0
                score -= 30.0 * abs(1.0 - aspect_ratio)
                if page_num > 0:
                    score -= 50.0 * page_num

                if rects:
                    r = rects[0]
                    score -= 50.0 * (r.y0 / page_h)
                    score -= 0.1 * abs(100.0 - r.width)

                candidates.append((score, image_bytes))

        if candidates and not is_scanned_doc:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]

        # =====================================================================
        # FALLBACK : Détection de visage avec OpenCV (si scan ou aucun candidat)
        # =====================================================================
        print("[INFO] Tentative de détection de visage (PDF scanné ou aucune image classique trouvée).")
        try:
            import cv2
            import numpy as np
            
            first_page = doc[0]
            # Zoom pour avoir une résolution correcte (2.0 = ~144 DPI)
            matrix = fitz.Matrix(2.0, 2.0)
            pix = first_page.get_pixmap(matrix=matrix)
            
            # Convertir Pixmap en array NumPy
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            
            # Gérer les canaux de couleur pour OpenCV (qui attend BGR ou BGRA)
            if pix.n == 4:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
            elif pix.n == 1:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)
                
            # Convertir en niveaux de gris pour la détection Haar Cascade
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Se concentrer sur le tiers supérieur de la page
            height, width = gray.shape
            upper_third_h = height // 3
            gray_upper_third = gray[:upper_third_h, :]
            
            # Charger le classificateur Haar Cascade
            cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
            face_cascade = cv2.CascadeClassifier(cascade_path)
            
            # Détection
            faces = face_cascade.detectMultiScale(
                gray_upper_third,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(30, 30)
            )
            
            if len(faces) > 0:
                # On prend le premier visage trouvé
                x, y, w, h = faces[0]
                
                # Ajouter une marge autour du visage (ex: 30%)
                margin = int(max(w, h) * 0.3)
                x_start = max(0, x - margin)
                y_start = max(0, y - margin)
                x_end = min(width, x + w + margin)
                y_end = min(upper_third_h, y + h + margin)
                
                # Recadrer l'image couleur originale (toujours RGB ici)
                face_img = img_array[y_start:y_end, x_start:x_end]
                
                # Convertir en BGR pour encoder correctement avec OpenCV
                face_img_bgr = cv2.cvtColor(face_img, cv2.COLOR_RGB2BGR)
                
                # Encoder en octets PNG
                success, buffer = cv2.imencode('.png', face_img_bgr)
                if success:
                    print("[INFO] Visage détecté et extrait avec succès par OpenCV.")
                    return buffer.tobytes()
            else:
                print("[INFO] Aucun visage détecté dans le tiers supérieur de la page.")
                
        except ImportError:
            print("[WARNING] OpenCV ou NumPy introuvables. Installez-les avec 'pip install opencv-python numpy'.")
        except Exception as ex:
            print(f"[WARNING] Erreur lors de la détection de visage : {ex}")

        # Dernier recours : renvoyer la meilleure image trouvée (s'il y en a) même si c'est un scan
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
            
        return None
    except Exception as e:
        print(f"Erreur d'extraction d'image : {e}")
        return None


def _extract_json_block(raw_text: str) -> str:
    """Extract first JSON object from model output, including fenced outputs."""
    cleaned = raw_text.strip()

    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", cleaned, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        return fenced.group(1)

    start = cleaned.find("{")
    if start == -1:
        raise ValueError("No JSON object found in model response.")

    depth = 0
    for idx in range(start, len(cleaned)):
        char = cleaned[idx]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return cleaned[start: idx + 1]

    raise ValueError("JSON object seems incomplete in model response.")


def _normalize_lmstudio_api_url(raw_url: str) -> str:
    """Force the OpenAI-compatible chat completions endpoint."""
    parsed = urlparse(raw_url)
    path = (parsed.path or "").rstrip("/")

    if path.endswith("/v1/chat/completions"):
        normalized_path = path
    elif path.endswith("/api/v1/chat"):
        normalized_path = path[: -len("/api/v1/chat")] + "/v1/chat/completions"
    elif path.endswith("/api/v1/chat/completions"):
        normalized_path = path[: -len("/api/v1/chat/completions")] + "/v1/chat/completions"
    else:
        normalized_path = "/v1/chat/completions"

    return urlunparse(parsed._replace(path=normalized_path))


def _build_chat_payload(cv_text: str, page_images_base64: Optional[List[str]] = None) -> Dict[str, Any]:
    schema = CVExtraction.model_json_schema()
    
    instruction = (
        "Tu es un expert en Ressources Humaines spécialisé dans l'analyse de Curriculum Vitae (CV).\n"
        "Ta mission est d'extraire les informations du CV fourni (texte ou image) de manière précise, exhaustive et structurée.\n\n"
        "### RÈGLES STRICTES :\n"
        "1. FIDÉLITÉ ABSOLUE : N'invente AUCUNE information. Base-toi exclusivement sur le contenu du CV.\n"
        "2. VALEURS MANQUANTES : Si une information est absente, laisse la liste vide `[]` ou la chaîne de caractères vide `\"\"` selon le type attendu. N'utilise pas de mentions comme 'Non spécifié'.\n"
        "3. CATÉGORISATION : Assure-toi de lister seulement les informations pour chaque clé,bien différencier les outils/technologies (Tools), les compétences techniques (Hard Skills) et les compétences comportementales (Soft Skills).\n"
        "4. FORMAT DE SORTIE : Réponds UNIQUEMENT avec un objet JSON valide. Ne rajoute aucun texte avant ou après, ni de balises markdown. Le JSON doit correspondre EXACTEMENT à ce schéma :\n"
        f"{json.dumps(schema, ensure_ascii=False, indent=2)}\n"
    )

    user_content: List[Dict[str, Any]] = []
    if cv_text.strip():
        user_content.append({"type": "text", "text": f"Texte extrait du CV:\n{cv_text}"})

    if page_images_base64:
        for image_b64 in page_images_base64:
            user_content.append(
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{image_b64}",
                        "detail": "high",
                    },
                }
            )

    if not user_content:
        user_content.append({"type": "text", "text": "Aucune donnée CV fournie."})

    return {
        "model": LMSTUDIO_MODEL,
        "messages": [
            {
                "role": "system",
                "content": instruction,
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
        "temperature": 0.1,
        "max_tokens": LMSTUDIO_MAX_TOKENS,
        "stream": False,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "cv_extraction",
                "schema": schema
            }
        }
    }


def _extract_chat_content(body: Any) -> str:
    if isinstance(body, dict):
        choices = body.get("choices")
        if isinstance(choices, list) and choices:
            first_choice = choices[0]
            if isinstance(first_choice, dict):
                finish_reason = first_choice.get("finish_reason")
                if finish_reason == "length":
                    raise RuntimeError(
                        "Réponse LM Studio tronquée avant la fin. "
                        "Augmentez LMSTUDIO_MAX_TOKENS ou LMSTUDIO_READ_TIMEOUT."
                    )

                message = first_choice.get("message")
                if isinstance(message, dict):
                    content = message.get("content", "")
                    if content:
                        return content
                    reasoning_content = message.get("reasoning_content", "")
                    if reasoning_content:
                        return reasoning_content

        message = body.get("message")
        if isinstance(message, dict):
            content = message.get("content", "")
            if content:
                return content
            reasoning_content = message.get("reasoning_content", "")
            if reasoning_content:
                return reasoning_content
        elif message is not None:
            return str(message)

        content = body.get("content")
        if content is not None:
            return content if isinstance(content, str) else json.dumps(content, ensure_ascii=False)

    return json.dumps(body, ensure_ascii=False)


def _parse_profile_json(content: str) -> Dict[str, Any]:
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError:
        json_text = _extract_json_block(content)
        parsed = json.loads(json_text)

    if not isinstance(parsed, dict):
        raise ValueError("Model output must be a JSON object.")

    return parsed

def extract_cv_data_vlm(cv_text: str, page_images_base64: Optional[List[str]] = None) -> str:
    """Call LM Studio VLM via OpenAI-compatible chat/completions endpoint."""
    payload = _build_chat_payload(cv_text=cv_text, page_images_base64=page_images_base64)
    lmstudio_url = _normalize_lmstudio_api_url(LMSTUDIO_API_URL)

    try:
        response = requests.post(
            lmstudio_url,
            json=payload,
            timeout=(LMSTUDIO_CONNECT_TIMEOUT, LMSTUDIO_READ_TIMEOUT),
            stream=False,
        )

        print("=" * 80)
        print("STATUS:", response.status_code)
        print("BODY:")
        print(response.text)
        print("=" * 80)

        response.raise_for_status()

    except ConnectionError as exc:
        raise RuntimeError(
            f"Impossible de joindre LM Studio sur {lmstudio_url}."
        ) from exc

    except Timeout as exc:
        raise RuntimeError(
            f"LM Studio n'a pas répondu dans le délai autorisé ({LMSTUDIO_CONNECT_TIMEOUT}s connexion, {LMSTUDIO_READ_TIMEOUT}s lecture)."
        ) from exc

    except requests.exceptions.HTTPError as exc:
        raise RuntimeError(
            f"HTTP Error {response.status_code}\n{response.text}"
        ) from exc

    try:
        body = response.json()
    except json.JSONDecodeError as exc:
        preview = response.text[:1000]
        raise RuntimeError(
            f"LM Studio a renvoyé une réponse JSON invalide: {preview}"
        ) from exc

    print(json.dumps(body, indent=2, ensure_ascii=False))

    content = _extract_chat_content(body)

    try:
        parsed = _parse_profile_json(content)
    except ValueError as exc:
        preview = content.strip().replace("\n", " ")[:1000]
        raise RuntimeError(
            f"Réponse LM Studio incomplète ou invalide. Aperçu: {preview}"
        ) from exc

    # Force la validation et le typage strict avec le modèle Pydantic Ezzahra
    try:
        validated = CVExtraction(**parsed).model_dump()
    except Exception as e:
        print(f"Attention, le JSON n'était pas parfait selon Pydantic : {e}")
        validated = parsed

    return json.dumps(validated, ensure_ascii=False)
