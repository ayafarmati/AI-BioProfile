import json
import urllib.request
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
import zipfile
from docx import Document
from pydantic import BaseModel, Field
from typing import List, Dict

# Définir le chemin par défaut de l'exécutable Tesseract sous Windows
TESSERACT_DEFAULT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(TESSERACT_DEFAULT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_DEFAULT_PATH

def perform_ocr_on_pdf(doc) -> str:
    """
    Parcourt les pages du document PDF, les convertit en images (pixmap)
    avec PyMuPDF, puis applique Tesseract OCR pour extraire le texte.
    """
    text = ""
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Utiliser un zoom de 150 DPI pour un bon compromis précision/rapidité
            zoom = 150 / 72.0
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix)
            
            # Convertir le pixmap en octets PNG
            img_bytes = pix.tobytes("png")
            
            # Charger dans PIL
            img = Image.open(io.BytesIO(img_bytes))
            
            # Appliquer l'OCR en français et anglais
            page_text = pytesseract.image_to_string(img, lang="fra+eng")
            text += f"--- Page {page_num + 1} ---\n"
            text += page_text + "\n"
            print(f"OCR de la page {page_num + 1}/{len(doc)} terminé.")
    except Exception as e:
        print(f"Erreur lors de l'exécution de l'OCR : {e}")
        # Si Tesseract n'est pas installé ou configuré, lever une erreur explicite
        if "tesseract is not installed or it's not in your PATH" in str(e).lower() or "[Errno 2]" in str(e):
            raise RuntimeError(
                "Tesseract OCR n'est pas détecté sur votre système. "
                "Veuillez l'installer depuis https://github.com/UB-Mannheim/tesseract/wiki "
                "et vous assurer que 'tesseract.exe' est présent dans 'C:\\Program Files\\Tesseract-OCR\\' "
                "ou dans votre variable d'environnement PATH."
            )
        else:
            raise e
    return text

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extrait le texte brut d'un fichier PDF à partir de ses octets.
    Si le document est scanné (pas de texte natif extrait), applique l'OCR local.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
        
    # Si le texte extrait est extrêmement court (seuil de 150 caractères), on applique l'OCR
    if len(text.strip()) < 150 and len(doc) > 0:
        print("[WARNING] PDF scanne detecte (texte extrait insuffisant). Lancement de l'OCR local...")
        text = perform_ocr_on_pdf(doc)
        
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extrait le texte brut d'un fichier DOCX à partir de ses octets.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Extraire aussi le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte du DOCX : {e}")
        return ""

def extract_image_from_pdf(file_bytes: bytes) -> bytes:
    """
    Parcourt les images du PDF et extrait la photo de profil probable en utilisant
    des critères de taille, d'aspect ratio et de position sur la page,
    tout en filtrant les fonds de page et les icônes.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
            
        candidates = []
        is_scanned_doc = False
        
        # Parcourir chaque page (surtout la première)
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_rect = page.rect
            page_w = page_rect.width
            page_h = page_rect.height
            image_list = page.get_images(full=True)
            
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                try:
                    # Utiliser extract_image qui gère automatiquement les masques de transparence (SMask)
                    extracted = doc.extract_image(xref)
                    raw_bytes = extracted["image"]
                    
                    # Utiliser PIL pour forcer un fond blanc si l'image a de la transparence
                    from PIL import Image
                    import io
                    img_pil = Image.open(io.BytesIO(raw_bytes))
                    if img_pil.mode in ('RGBA', 'LA') or (img_pil.mode == 'P' and 'transparency' in img_pil.info):
                        # Créer un fond blanc
                        bg = Image.new("RGB", img_pil.size, (255, 255, 255))
                        # Si palette avec transparence, convertir d'abord en RGBA
                        if img_pil.mode == 'P':
                            img_pil = img_pil.convert('RGBA')
                        # Obtenir le canal alpha pour le masque
                        alpha = img_pil.split()[-1]
                        bg.paste(img_pil, mask=alpha)
                        img_pil = bg
                    elif img_pil.mode != 'RGB':
                        img_pil = img_pil.convert('RGB')
                        
                    output = io.BytesIO()
                    img_pil.save(output, format="PNG")
                    image_bytes = output.getvalue()
                    
                    width = extracted["width"]
                    height = extracted["height"]
                except Exception as e:
                    print(f"[WARNING] Erreur lecture image xref {xref}: {e}")
                    continue
                
                # Aspect ratio en pixels
                aspect_ratio = width / height if height != 0 else 1.0
                if aspect_ratio < 0.4 or aspect_ratio > 2.5:
                    continue
                
                rects = page.get_image_rects(xref)
                
                if rects:
                    r = rects[0]
                    # Ignorer les arrière-plans ou pages complètes scannées
                    if r.width > page_w * 0.75 or r.height > page_h * 0.75:
                        is_scanned_doc = True
                        continue
                    if r.width > page_w * 0.5 and r.height > page_h * 0.5:
                        is_scanned_doc = True
                        continue
                    # Ignorer les petites icônes
                    if r.width < 25 or r.height < 25:
                        continue
                    # Ignorer les éléments tout en bas (comme les logos de pied de page)
                    if r.y0 > page_h * 0.7:
                        continue
                else:
                    # Fallback sur les dimensions pixel
                    if width < 100 or height < 100:
                        continue
                    if width > 900 or height > 900:
                        continue
                
                # Calculer un score de probabilité
                score = 100.0
                # Plus on s'éloigne d'un carré (1:1), plus on a une pénalité
                score -= 30.0 * abs(1.0 - aspect_ratio)
                
                # Favoriser la première page
                if page_num > 0:
                    score -= 50.0 * page_num
                    
                if rects:
                    r = rects[0]
                    # Favoriser le haut de la page
                    score -= 50.0 * (r.y0 / page_h)
                    # Favoriser une taille de profil idéale (~100 points de large)
                    score -= 0.1 * abs(100.0 - r.width)
                
                candidates.append((score, image_bytes))
                
        # Si la méthode classique a trouvé des candidats et que ce n'est pas un scan
        if candidates and not is_scanned_doc:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
            
        # =====================================================================
        # FALLBACK : Détection de visage avec OpenCV (si scan ou aucun candidat)
        # =====================================================================
        print("[INFO] Tentative de détection de visage (PDF scanné ou aucune image classique trouvée).")
        try:
            import cv2
            import numpy as np
            
            first_page = doc[0]
            # Zoom pour avoir une résolution correcte (2.0 = ~144 DPI)
            matrix = fitz.Matrix(2.0, 2.0)
            pix = first_page.get_pixmap(matrix=matrix)
            
            # Convertir Pixmap en array NumPy
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            
            # Gérer les canaux de couleur pour OpenCV (qui attend BGR ou BGRA)
            if pix.n == 4:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
            elif pix.n == 1:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)
                
            # Convertir en niveaux de gris pour la détection Haar Cascade
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Se concentrer sur le tiers supérieur de la page
            height, width = gray.shape
            upper_third_h = height // 3
            gray_upper_third = gray[:upper_third_h, :]
            
            # Charger le classificateur Haar Cascade
            cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
            face_cascade = cv2.CascadeClassifier(cascade_path)
            
            # Détection
            faces = face_cascade.detectMultiScale(
                gray_upper_third,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(30, 30)
            )
            
            if len(faces) > 0:
                # On prend le premier visage trouvé
                x, y, w, h = faces[0]
                
                # Ajouter une marge autour du visage (ex: 30%)
                margin = int(max(w, h) * 0.3)
                x_start = max(0, x - margin)
                y_start = max(0, y - margin)
                x_end = min(width, x + w + margin)
                y_end = min(upper_third_h, y + h + margin)
                
                # Recadrer l'image couleur originale (toujours RGB ici)
                face_img = img_array[y_start:y_end, x_start:x_end]
                
                # Convertir en BGR pour encoder correctement avec OpenCV
                face_img_bgr = cv2.cvtColor(face_img, cv2.COLOR_RGB2BGR)
                
                # Encoder en octets PNG
                success, buffer = cv2.imencode('.png', face_img_bgr)
                if success:
                    print("[INFO] Visage détecté et extrait avec succès par OpenCV.")
                    return buffer.tobytes()
            else:
                print("[INFO] Aucun visage détecté dans le tiers supérieur de la page.")
                
        except ImportError:
            print("[WARNING] OpenCV ou NumPy introuvables. Installez-les avec 'pip install opencv-python numpy'.")
        except Exception as ex:
            print(f"[WARNING] Erreur lors de la détection de visage : {ex}")

        # Dernier recours : renvoyer la meilleure image trouvée (s'il y en a) même si c'est un scan
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
            
        return None
    except Exception as e:
        print(f"Erreur d'extraction d'image : {e}")
        return None

def extract_image_from_docx(file_bytes: bytes) -> bytes:
    """
    Traite le fichier DOCX comme une archive ZIP et cherche l'image de profil
    dans le dossier word/media/.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            # Chercher tous les fichiers dans word/media/
            image_files = [f for f in z.namelist() if f.startswith('word/media/') and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
            
            if not image_files:
                return None
                
            candidates = []
            for img_path in image_files:
                with z.open(img_path) as img_f:
                    img_data = img_f.read()
                    
                    try:
                        from PIL import Image
                        img_pil = Image.open(io.BytesIO(img_data))
                        width, height = img_pil.size
                        
                        # Filtrer les petites icônes
                        if width < 50 or height < 50:
                            continue
                            
                        # Calculer un score basé sur le ratio (idéalement 1:1)
                        aspect_ratio = width / height if height != 0 else 1.0
                        score = 100.0 - (30.0 * abs(1.0 - aspect_ratio))
                        
                        # Favoriser les images de taille moyenne (photo de profil)
                        # au lieu d'énormes bannières de fond
                        if width > 1000 or height > 1000:
                            score -= 30.0
                            
                        # Gérer le fond transparent (remplacer par du blanc)
                        if img_pil.mode in ('RGBA', 'LA') or (img_pil.mode == 'P' and 'transparency' in img_pil.info):
                            bg = Image.new("RGB", img_pil.size, (255, 255, 255))
                            if img_pil.mode == 'P':
                                img_pil = img_pil.convert('RGBA')
                            alpha = img_pil.split()[-1]
                            bg.paste(img_pil, mask=alpha)
                            img_pil = bg
                        elif img_pil.mode != 'RGB':
                            img_pil = img_pil.convert('RGB')
                            
                        # Sauvegarder l'image modifiée
                        output = io.BytesIO()
                        img_pil.save(output, format="PNG")
                        img_data = output.getvalue()
                            
                        candidates.append((score, img_data))
                    except Exception:
                        continue
                        
            if candidates:
                # Trier par score décroissant
                candidates.sort(key=lambda x: x[0], reverse=True)
                return candidates[0][1]
                
        return None
    except Exception as e:
        print(f"Erreur d'extraction d'image depuis DOCX : {e}")
        return None

# ---------------------------------------------------------
# ---------------------------------------------------------
# CONFIGURATION DE L'API LOCALE DIRECTE (OLLAMA NATIVE)
# AVEC SCHEMA STRICT (Pydantic)
# ---------------------------------------------------------
LOCAL_OLLAMA_URL = "http://localhost:11434/api/chat" 
MODEL_NAME = "mistral"

class ProjetExperience(BaseModel):
    titre: str = Field(
        description="Le titre de l'expérience avec la date. Exemple: '02/2026 – Present: Ingénieur Architecture & Implantation véhicule'"
    )
    description: str = Field(
        description="Les tâches réalisées. Sépare chaque tâche par un retour à la ligne, SANS ajouter de puces (pas de tirets ni de points au début). Exemple: 'Analyser les exigences...\nConcevoir et développer...'"
    )

class CVExtraction(BaseModel):
    nom_complet: str = Field(
        description="Nom complet du candidat, idéalement tout en MAJUSCULES (ex: EZZAHRA ELMOUSSAOUY)."
    )
    titre_professionnel: str = Field(
        description="Le titre professionnel visé (ex: Ingénieur Architecture Implantation véhicule)."
    )
    disponibilite: str = Field(
        default="Immédiate",
        description="Disponibilité du candidat."
    )
    langues: List[str] = Field(
        description="Liste des langues. Exemple: ['Arabic (native)', 'French (Maitrise C2)']"
    )
    hard_skills: List[str] = Field(
        description="Liste des compétences techniques (Technical Skills). Exemple: ['Conception et modélisation 3D', 'Analyse de faisabilité']"
    )
    soft_skills: List[str] = Field(
        description="Liste des compétences comportementales (Soft Skills). Exemple: ['Esprit d'analyse', 'Rigueur technique']"
    )
    outils_et_technologies: List[str] = Field(
        description="Liste des outils logiciels (Tools). Exemple: ['CATIA V6', 'PLM']"
    )
    autres_informations: str = Field(
        default="",
        description="Toutes les autres informations présentes sur le CV (email, téléphone, adresse, liens, loisirs, etc.) qui ne correspondent pas EXACTEMENT aux autres champs. Mets TOUT le reste ici pour ne pas polluer les outils ou compétences."
    )
    projets_et_experiences: List[ProjetExperience] = Field(
        description="Liste des expériences professionnelles et projets pertinents."
    )

def extract_cv_data(cv_text: str) -> str:
    import json
    import urllib.request
    print(f"Envoi du CV au modèle {MODEL_NAME} en direct (Format JSON Strict via Pydantic)...")

    schema_json = CVExtraction.model_json_schema()

    system_prompt = (
        "Tu es un expert RH en extraction de données de CV.\n"
        "Extrais toutes les informations du CV fourni.\n"
        "Ne génère STRICTEMENT RIEN d'inventé, base-toi uniquement sur le texte du CV.\n"
        "Si une information n'est pas mentionnée, laisse la liste vide ou la chaîne vide.\n"
        "\nExemple de comportement attendu :\n"
        "Si le CV contient 'Développeur Python. Email: jean@mail.com. Outils: Git, Docker.', tu dois mettre 'Git' et 'Docker' dans outils_et_technologies, et 'Email: jean@mail.com' dans autres_informations."
    )

    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": cv_text}
        ],
        "format": schema_json,
        "stream": False
    }

    req = urllib.request.Request(
        LOCAL_OLLAMA_URL, 
        data=json.dumps(data).encode('utf-8'),
        headers={'Content-Type': 'application/json'}
    )

    try:
        with urllib.request.urlopen(req) as response:
            result = json.loads(response.read().decode('utf-8'))
            content = result['message']['content']
            return content
    except Exception as e:
        raise Exception(f"Erreur de communication avec Ollama: {e}")

if __name__ == "__main__":
    import sys
    import os
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = "cv_test.pdf" # Default test if none provided
        
    if not os.path.exists(file_path):
        print(f"Erreur : Le fichier {file_path} n'existe pas.")
        sys.exit(1)
        
    print(f"Lecture du fichier : {file_path}")
    
    if file_path.lower().endswith(".pdf"):
        with open(file_path, "rb") as f:
            cv_content = extract_text_from_pdf(f.read())
    elif file_path.lower().endswith(".docx"):
        with open(file_path, "rb") as f:
            cv_content = extract_text_from_docx(f.read())
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            cv_content = f.read()
    
    try:
        json_result = extract_cv_data(cv_content)
        print("\n=== RÉSULTAT DE L'EXTRACTION JSON STRICT ===")
        print(json_result)
        
        with open("resultat_cv.json", "w", encoding="utf-8") as out_f:
            out_f.write(json_result)
        print("\n[SUCCESS] Le resultat a ete sauvegarde dans 'resultat_cv.json'")
    except Exception as e:
        print(f"\nErreur : {e}")
