import os
import json
import urllib.request
import urllib.error
from typing import List, Optional, Any, Union, Dict

import fitz  # PyMuPDF
import pymupdf4llm
import pytesseract
from pydantic import BaseModel, Field
import io
import zipfile
from docx import Document
from dotenv import load_dotenv

load_dotenv()

DEFAULT_OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:3b")
DEFAULT_OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434/api/chat")
# ==============================================================================
# CONFIGURATION TESSERACT
# ==============================================================================
TESSERACT_DEFAULT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(TESSERACT_DEFAULT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_DEFAULT_PATH

# ==============================================================================
# MODÈLES PYDANTIC 
# ==============================================================================

class Experience(BaseModel):
    titre_experience: str = Field(
        default="Expérience ou Projet",
        description="La fonction officielle, le rôle ou le poste occupé par le candidat lors d'une expérience professionnelle spécifique, OU bien le titre d'un projet (académique, personnel, etc.). Mots-clés / Exemples : '02/2026 – Present: Ingénieur Architecture', 'Projet de fin d\\'études : Robot Autonome'."
    )
    details_experience: List[str] = Field(
        default_factory=list,
        description="RÈGLE ABSOLUE : Copier-coller EXACTEMENT les phrases (bullet points) décrivant les tâches depuis le texte source. Ne RIEN résumer, ne RIEN reformuler. Chaque élément de la liste correspond à une puce/bullet point."
    )

class CVExtraction(BaseModel):
    analyse_preliminaire: str = Field(
        default="",
        description="Rédige ici, EN FRANÇAIS ou ANGLAIS, une analyse préliminaire globale du CV. C'est ton brouillon de pensée (Chain of Thought). Fais le tri mentalement entre Outils et Compétences avant de remplir le reste."
    )
    nom_complet: str = Field(
        default="Candidat Inconnu",
        description="Le prénom et le nom de famille du candidat. Mots-clés / Exemples : Jean Dupont, Ezzahra Elmoussaouy."
    )
    disponibilite: str = Field(
        default="Immédiate",
        description="Disponibilité du candidat."
    )
    email: str = Field(
        default="",
        description="L'adresse email du candidat. Laisser vide si non trouvée."
    )
    telephone: str = Field(
        default="",
        description="Le numéro de téléphone du candidat. Laisser vide si non trouvé."
    )
    liens: List[Any] = Field(
        default_factory=list,
        description="Les liens web (LinkedIn, GitHub, Portfolio, etc.)."
    )
    titre_professionnel: str = Field(
        default="Profil Général",
        description="Le titre global ou l'en-tête du profil du candidat, définissant son métier principal actuel ou visé. Ne pas confondre avec les titres de postes spécifiques. Mots-clés / Exemples : 'Ingénieur Architecture Implantation véhicule', 'Développeur Fullstack'."
    )
    tools: List[Any] = Field(
        default_factory=list,
        description="Les logiciels, langages de programmation, plateformes, frameworks ou équipements utilisés pour appliquer une compétence technique. (Règle : si ça s'installe, se télécharge ou s'achète, c'est un outil). Mots-clés / Exemples : CATIA V6, Python, Microsoft Excel, Jira, React.js."
    )
    soft_skills: List[Any] = Field(
        default_factory=list,
        description="Les compétences comportementales (Soft Skills). RÈGLE ABSOLUE : Tu ne dois extraire que les soft skills EXPLICITEMENT ÉCRITES dans le CV. IL EST STRICTEMENT INTERDIT d'inventer, de déduire ou de deviner des soft skills. Si le CV ne mentionne aucune soft skill textuellement, retourne une liste VIDE []."
    )
    technical_skills: List[Any] = Field(
        default_factory=list,
        description="Les compétences techniques pures (concepts, méthodes). RÈGLE ABSOLUE : IL EST STRICTEMENT INTERDIT de copier-coller des phrases entières ou des descriptions de tâches provenant de la section 'Expériences'. Tu dois uniquement extraire les compétences listées dans les rubriques de compétences. Si aucune n'est listée, retourne une liste VIDE []."
    )
    langues: List[Any] = Field(
        default_factory=list,
        description="Les langues parlées ou écrites par le candidat, accompagnées de leur niveau de maîtrise. Mots-clés / Exemples : Arabe (Natif), Français (Maîtrise C2), Anglais (Courant)."
    )
    formations: List[Any] = Field(
        default_factory=list,
        description="Les formations, diplômes, études ou cursus académiques suivis par le candidat. Champ optionnel, à extraire uniquement si présent."
    )
    experiences: List[Experience] = Field(
        default_factory=list,
        description="La liste complète des expériences professionnelles, stages ET TOUS LES PROJETS (académiques, personnels, etc.). Tout projet doit être extrait ici comme une expérience."
    )
    autres_informations: str = Field(
        default="",
        description="RÈGLE CRITIQUE : Tu dois copier-coller ici TOUTES les autres rubriques du CV qui ne rentrent pas dans les champs précédents (par exemple : Certifications, Centres d'intérêt, Bénévolat...). Garde les titres originaux du CV sous forme de texte clair avec des retours à la ligne."
    )


# ==============================================================================
# LOGIQUE D'EXTRACTION DE TEXTE
# ==============================================================================



def extract_text_from_pdf_as_markdown(file_bytes: bytes) -> str:
    """
    Extrait le texte d'un PDF. Repose sur le moteur natif de pymupdf4llm 
    (Intelligence Artificielle) pour détecter automatiquement l'ordre de lecture 
    des colonnes et de la mise en page.
    """
    print("[INFO] Lancement de l'extraction de texte (Markdown natif)...")
    try:
        doc = fitz.Document(stream=file_bytes, filetype="pdf")
        
        # Concaténation rapide du texte brut pour vérifier si le PDF est natif
        text_length = 0
        for page in doc:
            text_length += len(page.get_text())
            if text_length >= 50:
                break
                
        if text_length >= 50:
            print("[INFO] PDF natif détecté. Analyse avec le layout ML de pymupdf4llm en cours...")
            md_text = pymupdf4llm.to_markdown(doc)
            doc.close()
            return md_text
            
        print("[INFO] PDF scanné ou sans texte détecté (<50 caractères). Basculement vers Tesseract OCR pleine page...")
        extracted_text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Zoom pour améliorer la qualité de l'OCR
            zoom = 2.0 
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # Conversion de l'image (pixmap) en format compatible PIL/Tesseract
            from PIL import Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            print(f"[INFO] Page {page_num+1} (Scannée) : Extraction OCR pleine page.")
            page_text = pytesseract.image_to_string(img, lang="fra+eng")
            extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
        doc.close()
        return extracted_text

    except Exception as e:
        print(f"[ERREUR] Erreur lors de l'extraction du texte PDF : {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extrait le texte d'un fichier DOCX en le formatant en Markdown 
    pour préserver la structure (titres, puces, tableaux).
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            style_name = para.style.name.lower()
            
            # Détection des titres pour le Markdown
            if 'heading 1' in style_name or 'titre 1' in style_name:
                lines.append(f"\n# {text}")
            elif 'heading 2' in style_name or 'titre 2' in style_name:
                lines.append(f"\n## {text}")
            elif 'heading 3' in style_name or 'titre 3' in style_name:
                lines.append(f"\n### {text}")
            # Détection des listes à puces
            elif 'list' in style_name or 'puce' in style_name:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        
        # Extraction et formatage des tableaux en Markdown
        if doc.tables:
            lines.append("\n### Tableaux ou dispositions en colonnes :")
            for table in doc.tables:
                lines.append("")
                for row_idx, row in enumerate(table.rows):
                    # Remplacer les retours à la ligne dans les cellules par des espaces pour le Markdown
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    lines.append("| " + " | ".join(row_data) + " |")
                    
                    # Ajouter le séparateur markdown après la première ligne
                    if row_idx == 0:
                        separator = ["---"] * len(row.cells)
                        lines.append("| " + " | ".join(separator) + " |")
                lines.append("")
        
        return "\n".join(lines)
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte du DOCX : {e}")
        return ""

def extract_image_from_pdf(file_bytes: bytes) -> bytes:
    """
    Parcourt les images du PDF et extrait la photo de profil probable en utilisant
    des critères de taille, d'aspect ratio et de position sur la page,
    tout en filtrant les fonds de page et les icônes.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
            
        candidates = []
        is_scanned_doc = False
        
        # Parcourir chaque page (surtout la première)
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_rect = page.rect
            page_w = page_rect.width
            page_h = page_rect.height
            image_list = page.get_images(full=True)
            
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                try:
                    # Utiliser extract_image qui gère automatiquement les masques de transparence (SMask)
                    extracted = doc.extract_image(xref)
                    raw_bytes = extracted["image"]
                    
                    # Utiliser PIL pour forcer un fond blanc si l'image a de la transparence
                    from PIL import Image
                    import io
                    img_pil = Image.open(io.BytesIO(raw_bytes))
                    if img_pil.mode in ('RGBA', 'LA') or (img_pil.mode == 'P' and 'transparency' in img_pil.info):
                        # Créer un fond blanc
                        bg = Image.new("RGB", img_pil.size, (255, 255, 255))
                        # Si palette avec transparence, convertir d'abord en RGBA
                        if img_pil.mode == 'P':
                            img_pil = img_pil.convert('RGBA')
                        # Obtenir le canal alpha pour le masque
                        alpha = img_pil.split()[-1]
                        bg.paste(img_pil, mask=alpha)
                        img_pil = bg
                    elif img_pil.mode != 'RGB':
                        img_pil = img_pil.convert('RGB')
                        
                    output = io.BytesIO()
                    img_pil.save(output, format="PNG")
                    image_bytes = output.getvalue()
                    
                    width = extracted["width"]
                    height = extracted["height"]
                except Exception as e:
                    print(f"[WARNING] Erreur lecture image xref {xref}: {e}")
                    continue
                
                # Aspect ratio en pixels
                aspect_ratio = width / height if height != 0 else 1.0
                if aspect_ratio < 0.4 or aspect_ratio > 2.5:
                    continue
                
                rects = page.get_image_rects(xref)
                
                if rects:
                    r = rects[0]
                    # Ignorer les arrière-plans ou pages complètes scannées
                    if r.width > page_w * 0.75 or r.height > page_h * 0.75:
                        is_scanned_doc = True
                        continue
                    if r.width > page_w * 0.5 and r.height > page_h * 0.5:
                        is_scanned_doc = True
                        continue
                    # Ignorer les petites icônes
                    if r.width < 25 or r.height < 25:
                        continue
                    # Ignorer les éléments tout en bas (comme les logos de pied de page)
                    if r.y0 > page_h * 0.7:
                        continue
                else:
                    # Fallback sur les dimensions pixel
                    if width < 100 or height < 100:
                        continue
                    if width > 900 or height > 900:
                        continue
                
                # Calculer un score de probabilité
                score = 100.0
                # Plus on s'éloigne d'un carré (1:1), plus on a une pénalité
                score -= 30.0 * abs(1.0 - aspect_ratio)
                
                # Favoriser la première page
                if page_num > 0:
                    score -= 50.0 * page_num
                    
                if rects:
                    r = rects[0]
                    # Favoriser le haut de la page
                    score -= 50.0 * (r.y0 / page_h)
                    # Favoriser une taille de profil idéale (~100 points de large)
                    score -= 0.1 * abs(100.0 - r.width)
                
                candidates.append((score, image_bytes))
                
        # Si la méthode classique a trouvé des candidats et que ce n'est pas un scan
        if candidates and not is_scanned_doc:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
            
        # =====================================================================
        # FALLBACK : Détection de visage avec OpenCV (si scan ou aucun candidat)
        # =====================================================================
        print("[INFO] Tentative de détection de visage (PDF scanné ou aucune image classique trouvée).")
        try:
            import cv2
            import numpy as np
            
            first_page = doc[0]
            # Zoom pour avoir une résolution correcte (2.0 = ~144 DPI)
            matrix = fitz.Matrix(2.0, 2.0)
            pix = first_page.get_pixmap(matrix=matrix)
            
            # Convertir Pixmap en array NumPy
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            
            # Gérer les canaux de couleur pour OpenCV (qui attend BGR ou BGRA)
            if pix.n == 4:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
            elif pix.n == 1:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)
                
            # Convertir en niveaux de gris pour la détection Haar Cascade
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Se concentrer sur le tiers supérieur de la page
            height, width = gray.shape
            upper_third_h = height // 3
            gray_upper_third = gray[:upper_third_h, :]
            
            # Charger le classificateur Haar Cascade
            cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
            face_cascade = cv2.CascadeClassifier(cascade_path)
            
            # Détection
            faces = face_cascade.detectMultiScale(
                gray_upper_third,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(30, 30)
            )
            
            if len(faces) > 0:
                # On prend le premier visage trouvé
                x, y, w, h = faces[0]
                
                # Ajouter une marge autour du visage (ex: 30%)
                margin = int(max(w, h) * 0.3)
                x_start = max(0, x - margin)
                y_start = max(0, y - margin)
                x_end = min(width, x + w + margin)
                y_end = min(upper_third_h, y + h + margin)
                
                # Recadrer l'image couleur originale (toujours RGB ici)
                face_img = img_array[y_start:y_end, x_start:x_end]
                
                # Convertir en BGR pour encoder correctement avec OpenCV
                face_img_bgr = cv2.cvtColor(face_img, cv2.COLOR_RGB2BGR)
                
                # Encoder en octets PNG
                success, buffer = cv2.imencode('.png', face_img_bgr)
                if success:
                    print("[INFO] Visage détecté et extrait avec succès par OpenCV.")
                    return buffer.tobytes()
            else:
                print("[INFO] Aucun visage détecté dans le tiers supérieur de la page.")
                
        except ImportError:
            print("[WARNING] OpenCV ou NumPy introuvables. Installez-les avec 'pip install opencv-python numpy'.")
        except Exception as ex:
            print(f"[WARNING] Erreur lors de la détection de visage : {ex}")

        # Dernier recours : renvoyer la meilleure image trouvée (s'il y en a) même si c'est un scan
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
            
        return None
    except Exception as e:
        print(f"Erreur d'extraction d'image : {e}")
        return None

def extract_image_from_docx(file_bytes: bytes) -> bytes:
    """
    Traite le fichier DOCX comme une archive ZIP et cherche l'image de profil
    dans le dossier word/media/.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            # Chercher tous les fichiers dans word/media/
            image_files = [f for f in z.namelist() if f.startswith('word/media/') and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
            
            if not image_files:
                return None
                
            candidates = []
            for img_path in image_files:
                with z.open(img_path) as img_f:
                    img_data = img_f.read()
                    
                    try:
                        from PIL import Image
                        img_pil = Image.open(io.BytesIO(img_data))
                        width, height = img_pil.size
                        
                        # Filtrer les petites icônes
                        if width < 50 or height < 50:
                            continue
                            
                        # Calculer un score basé sur le ratio (idéalement 1:1)
                        aspect_ratio = width / height if height != 0 else 1.0
                        score = 100.0 - (30.0 * abs(1.0 - aspect_ratio))
                        
                        # Favoriser les images de taille moyenne (photo de profil)
                        # au lieu d'énormes bannières de fond
                        if width > 1000 or height > 1000:
                            score -= 30.0
                            
                        # Gérer le fond transparent (remplacer par du blanc)
                        if img_pil.mode in ('RGBA', 'LA') or (img_pil.mode == 'P' and 'transparency' in img_pil.info):
                            bg = Image.new("RGB", img_pil.size, (255, 255, 255))
                            if img_pil.mode == 'P':
                                img_pil = img_pil.convert('RGBA')
                            alpha = img_pil.split()[-1]
                            bg.paste(img_pil, mask=alpha)
                            img_pil = bg
                        elif img_pil.mode != 'RGB':
                            img_pil = img_pil.convert('RGB')
                            
                        # Sauvegarder l'image modifiée
                        output = io.BytesIO()
                        img_pil.save(output, format="PNG")
                        img_data = output.getvalue()
                            
                        candidates.append((score, img_data))
                    except Exception:
                        continue
                        
            if candidates:
                # Trier par score décroissant
                candidates.sort(key=lambda x: x[0], reverse=True)
                return candidates[0][1]
                
        return None
    except Exception as e:
        print(f"Erreur d'extraction d'image depuis DOCX : {e}")
        return None

# ==============================================================================
# COMMUNICATION AVEC OLLAMA (VIA URLLIB)
# ==============================================================================

def _build_ollama_payload(cv_markdown: str, model_name: str) -> dict:
    """
    Construit le payload (corps de la requête) pour l'API /api/chat d'Ollama.
    """
    # Template JSON ultra-clair pour éviter que le modèle ne s'emmêle avec les $refs de Pydantic
    json_template = """{
  "nom_complet": "Prénom NOM",
  "email": "email@example.com",
  "telephone": "+212 600 000 000",
  "liens": ["https://linkedin.com/in/...", "https://github.com/..."],
  "disponibilite": "Immédiate",
  "titre_professionnel": "Le titre global du profil",
  "tools": ["Nom de logiciel 1", "Langage 2"],
  "soft_skills": ["Soft skill 1", "Soft skill 2"],
  "technical_skills": ["Concept technique 1", "Savoir-faire 2"],
  "formations": ["Diplôme 1 (École, Année)", "Formation 2"],
  "langues": ["Langue 1 (Niveau)", "Langue 2 (Niveau)"],
  "experiences": [
    {
      "titre_experience": "Titre du 1er poste | Dates",
      "details_experience": [
        "Tâche accomplie 1...",
        "Tâche accomplie 2..."
      ]
    },
    {
      "titre_experience": "Titre du 2ème poste ou projet | Dates",
      "details_experience": [
        "Tâche accomplie..."
      ]
    }
  ],
  "autres_informations": "Certifications, Centres d'intérêt, Bénévolat..."
}"""
    
    
    system_prompt = f"""Contexte et Rôle :
Tu es un simple outil d'extraction JSON (rôle basique). Ton seul objectif est de copier-coller les informations du texte source vers le bon champ JSON. IL EST STRICTEMENT INTERDIT de reformuler, résumer, traduire ou inventer du texte.

Définition des zones d'extraction :

1. Nom complet (Full Name)
Définition : Le prénom et le nom de famille du candidat.
Mots-clés / Exemples : Jean Dupont, Ezzahra Elmoussaouy.

2. Titre professionnel (Profile Title)
Définition : Le titre global ou l'en-tête du profil du candidat, définissant son métier principal actuel ou visé. Ne pas confondre avec les titres de postes spécifiques dans la section expérience.
Mots-clés / Exemples : "Ingénieur Architecture Implantation véhicule", "Développeur Fullstack Senior", "Consultant Data".

2b. Contact (Email, Téléphone, Liens)
Définition : Extrais l'adresse email, le numéro de téléphone, et liste tous les hyperliens (URL) trouvés dans le CV (LinkedIn, GitHub, Portfolio). Si l'un est manquant, laisse le champ vide ou une liste vide [].

3. Expériences et Projets (Experiences)
Définition : C'est une liste (tableau) de TOUTES les expériences professionnelles, stages ET de TOUS les projets (académiques, professionnels, personnels). Chaque projet DOIT être extrait comme une expérience à part entière. Chaque élément de la liste doit contenir deux champs obligatoires :
  - 'titre_experience' : La fonction, le rôle ou le titre du projet, incluant l'entreprise ou le contexte (Ex: "02/2026 – Present: Ingénieur Architecture", "Projet Académique : Robot Autonome").
  - 'details_experience' : Un tableau listant EXACTEMENT les puces (bullet points) accomplies. Tu dois COPIER-COLLER le texte source mot pour mot. Chaque puce du document original doit devenir un élément de cette liste. Ne fais aucun résumé.
ATTENTION : N'oublie aucune expérience ni AUCUN PROJET mentionné dans le document ! S'il y a une section 'Projets', ils vont tous ici.
Si un fragment de texte isolé ne semble appartenir à aucune expérience ou projet du corps principal, ignore-le simplement.

4. Technical Skills (Compétences Techniques / Savoir-faire)
Définition : Les domaines d'expertise, concepts théoriques, méthodologies métier et savoir-faire pratiques. C'est la méthode de travail ou le concept (indépendant d'un logiciel spécifique).
Mots-clés / Exemples : Développement Web, Modélisation & Conception, Machine Learning, Bases de Données, Analyse de données, Gestion de projet.
ATTENTION : Tu dois être strictement EXHAUSTIF. Extrais TOUS les domaines d'expertise listés dans les rubriques de compétences ET les savoir-faire mentionnés dans les descriptions de projets (ex: conception d'un modèle de prédiction). N'y insère AUCUN soft skill (comme Autonomie, Travail en équipe).

5. Tools (Outils / Moyens techniques)
Définition : Les logiciels, langages de programmation, plateformes, frameworks ou équipements utilisés pour appliquer une compétence technique. (Règle : si ça s'installe, se télécharge ou s'achète, c'est un outil).
Mots-clés / Exemples : CATIA V6, Python, Microsoft Excel, Jira, React.js.
ATTENTION : Tu dois être strictement EXHAUSTIF. Extrais TOUS les outils mentionnés dans le document sans aucune exception. Ne te limite surtout pas aux premiers que tu trouves.

6. Soft Skills (Compétences Comportementales / Savoir-être)
Définition : Les qualités humaines, relationnelles, émotionnelles et d'organisation du candidat. Elles décrivent comment la personne se comporte au travail.
Mots-clés / Exemples : Esprit d'analyse, Travail en équipe, Adaptabilité, Leadership, Communication, Résolution de problèmes, Autonomie.
RÈGLE ABSOLUE : Un soft skill est TOUJOURS un mot unique ou un groupe de 2 à 3 mots maximum (un nom ou un adjectif). Tu ne dois JAMAIS extraire de phrases complètes ou de verbes conjugués. Si le candidat ne liste pas explicitement de mots-clés de savoir-être, retourne une liste VIDE [].

7. Formations (Education)
Définition : Les diplômes, études supérieures, écoles ou universités fréquentées par le candidat.
Mots-clés / Exemples : "Diplôme d'Ingénieur d'État en Informatique (2020 - 2023)", "Master en Data Science".
ATTENTION : Ce champ est optionnel. Si aucune formation n'est mentionnée de manière claire, retourne une liste VIDE [].

8. Langues (Languages)
Définition : Les langues parlées ou écrites par le candidat, accompagnées de leur niveau de maîtrise.
Mots-clés / Exemples : Arabe (Natif), Français (Maîtrise C2), Anglais (Courant).

9. Autres informations (Bloc de récupération global)
Définition : C'est le champ le plus important pour ne perdre aucune donnée. Tu DOIS y insérer, sous forme de texte brut avec des retours à la ligne, TOUTES les rubriques du CV qui n'ont pas leur place précise dans les listes précédentes.
INCLUSIONS OBLIGATOIRES :
- La hiérarchie exacte des "Compétences Techniques" avec leurs sous-catégories (ex: "CAO : SolidWorks, CATIA", "Programmation : Python, C", "Embarqué : Arduino").
- Les "Certifications" complètes.
- Les "Centres d'Intérêt" et "Engagements associatifs".
- Les Coordonnées complètes (Adresse physique, Téléphone, Email, Liens GitHub/LinkedIn).
RÈGLE ABSOLUE : Préserve les titres du CV dans ce champ (ex: "Centres d'intérêt : Voyages..."). Ne résume rien, copie le texte.

Instructions de sortie :
Analyse le texte fourni et retourne les informations extraites uniquement sous le format JSON tel que défini par le schéma Pydantic. Remplis d'abord le champ d'analyse préliminaire, puis extrais les outils, puis le reste. Si une information est absente, laisse la liste vide [].
RÈGLE ABSOLUE D'EXHAUSTIVITÉ : Pour les listes (outils, compétences, langues, expériences), tu dois extraire TOUT ce qui est présent dans le CV. N'oublie aucun élément.
CORRECTION OCR : Le texte brut peut contenir des fautes de frappe ou d'orthographe car il provient parfois d'un scanner (ex: 'Outils' écrit 'Qutils', lettres accentuées cassées). Corrige-les naturellement en français lors de l'extraction de l'information.
Le résultat final DOIT être un document JSON valide et rien d'autre. Pas de blabla, ni d'explications.

Voici le modèle exact du JSON que tu DOIS renvoyer (remplis les valeurs avec les données extraites) :
{json_template}"""

    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Voici le texte structuré du CV à analyser :\n\n{cv_markdown}"}
        ],
        "format": "json",     # On force la sortie JSON native d'Ollama
        "stream": False,      # On veut toute la réponse d'un coup (pas de streaming)
        "options": {
            "temperature": 0.1, # Température basse pour privilégier la précision
            "num_ctx": 8192,    # Agrandir drastiquement la fenêtre de contexte (évite la saturation)
            "num_predict": -1   # Forcer le modèle à écrire jusqu'au bout, sans couper la fin
        }
    }
    return payload


def _call_ollama_api_urllib(payload: dict, url: str) -> dict:
    """
    Appelle l'API locale Ollama en utilisant exclusivement la librairie standard `urllib`.
    Gère les erreurs HTTP, Timeout et problèmes de décodage JSON.
    """
    # Encodage du payload en bytes pour l'envoi HTTP POST
    data = json.dumps(payload).encode('utf-8')
    
    # Préparation de la requête avec les bons en-têtes
    req = urllib.request.Request(url, data=data, method="POST")
    req.add_header('Content-Type', 'application/json')
    
    try:
        # Appel API avec un timeout généreux (les LLMs locaux peuvent être lents)
        print(f"[INFO] Appel API Ollama vers {url} en cours...")
        with urllib.request.urlopen(req, timeout=1800) as response:
            response_data = response.read().decode('utf-8')
            
            try:
                # Analyse de la réponse d'Ollama (JSON)
                json_response = json.loads(response_data)
                return json_response
            except json.JSONDecodeError as e:
                raise ValueError(f"La réponse de l'API Ollama n'est pas un JSON valide : {e}\nRéponse: {response_data}")
                
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"Erreur HTTP serveur Ollama : {e.code} - {e.reason}")
    except urllib.error.URLError as e:
        raise RuntimeError(f"Impossible de se connecter à Ollama à l'adresse {url}. Assurez-vous qu'Ollama est lancé. Erreur : {e.reason}")
    except TimeoutError:
        raise RuntimeError(f"L'appel à l'API Ollama a dépassé le délai imparti (Timeout).")
    except Exception as e:
        raise RuntimeError(f"Erreur inattendue lors de l'appel à Ollama : {e}")


# ==============================================================================
# FONCTION PRINCIPALE (ORCHESTRATION)
# ==============================================================================

def extract_cv_data_llm(file_bytes: bytes, model_name: str = None, ollama_url: str = None) -> str:
    """
    Orchestre le processus d'extraction complet :
    1. PDF -> Markdown (ou OCR Tesseract)
    2. Construction du Prompt avec le Schéma Pydantic
    3. Appel à Ollama via urllib
    4. Validation de la réponse brute via Pydantic
    5. Retour du JSON validé sous forme de chaîne.
    """
    model_name = model_name or DEFAULT_OLLAMA_MODEL
    ollama_url = ollama_url or DEFAULT_OLLAMA_URL
    
    # 1. Extraction en Markdown ou OCR
    print("[INFO] Démarrage de l'extraction de texte...")
    cv_markdown = extract_text_from_pdf_as_markdown(file_bytes)
    
    if not cv_markdown.strip():
        raise ValueError("Aucun texte n'a pu être extrait du PDF.")
        
    print(f"[INFO] Extraction réussie ({len(cv_markdown)} caractères).")
    
    # Sauvegarde pour debug/visualisation (comme dans test_extraction_v2)
    with open("debug_extracted_text.txt", "w", encoding="utf-8") as f:
        f.write(cv_markdown)
    print("[DEBUG] Le texte Markdown/OCR envoyé au LLM a été sauvegardé dans 'debug_extracted_text.txt'")
    
    # 2. Construction du payload
    print(f"[INFO] Préparation de la requête pour Ollama (modèle: {model_name})...")
    payload = _build_ollama_payload(cv_markdown, model_name)
    
    # 3. Appel de l'API (utilisation exclusive de urllib)
    api_response = _call_ollama_api_urllib(payload, ollama_url)
    
    # Extraction du message généré par l'assistant
    try:
        raw_json_str = api_response["message"]["content"]
    except KeyError:
        raise ValueError("La structure de la réponse d'Ollama est invalide (clé 'message' ou 'content' manquante).")
    
    # 4. Validation via Pydantic
    print("[INFO] Validation des données extraites avec Pydantic...")
    try:
        # On essaie de charger le JSON brute répondu par l'Ollama
        parsed_dict = json.loads(raw_json_str)
        
        # APLATISSEMENT DES LISTES
        for field in ["formations", "langues", "tools", "soft_skills", "technical_skills", "liens"]:
            if field in parsed_dict and isinstance(parsed_dict[field], list):
                new_list = []
                for item in parsed_dict[field]:
                    if isinstance(item, str):
                        new_list.append(item)
                    elif isinstance(item, dict):
                        new_list.append(" - ".join(str(v) for v in item.values() if v))
                    else:
                        new_list.append(str(item))
                parsed_dict[field] = new_list
                
        # Validation stricte et instanciation du modèle Pydantic
        validated_data = CVExtraction(**parsed_dict)
        
        # 5. On renvoie le JSON propre et validé sous forme de string
        final_json_str = json.dumps(validated_data.model_dump(), ensure_ascii=False, indent=4)
        print("[INFO] Extraction et validation terminées avec succès.")
        return final_json_str
        
    except json.JSONDecodeError as e:
        print(f"[ERREUR] Le modèle Ollama n'a pas renvoyé un JSON valide: {raw_json_str}")
        raise ValueError(f"JSON invalide renvoyé par le modèle: {e}")
    except Exception as e:
        print(f"[ERREUR] Échec de la validation Pydantic : {e}")
        raise ValueError(f"Le JSON renvoyé ne correspond pas au schéma attendu : {e}")


def extract_cv_data_llm_from_text(cv_text: str, model_name: str = None, ollama_url: str = None) -> str:
    """
    Orchestre le processus d'extraction complet à partir de texte brut (utile pour DOCX et TXT).
    """
    model_name = model_name or DEFAULT_OLLAMA_MODEL
    ollama_url = ollama_url or DEFAULT_OLLAMA_URL
    
    if not cv_text.strip():
        raise ValueError("Aucun texte fourni.")
        
    # Sauvegarde pour debug/visualisation
    with open("debug_extracted_text.txt", "w", encoding="utf-8") as f:
        f.write(cv_text)
    print("[DEBUG] Le texte brut envoyé au LLM a été sauvegardé dans 'debug_extracted_text.txt'")
        
    print(f"[INFO] Préparation de la requête pour Ollama à partir de texte brut (modèle: {model_name})...")
    payload = _build_ollama_payload(cv_text, model_name)
    api_response = _call_ollama_api_urllib(payload, ollama_url)
    
    try:
        raw_json_str = api_response["message"]["content"]
    except KeyError:
        raise ValueError("La structure de la réponse d'Ollama est invalide.")
    
    print("[INFO] Validation des données extraites avec Pydantic...")
    try:
        parsed_dict = json.loads(raw_json_str)
        
        # APLATISSEMENT DES LISTES : On s'assure que les champs censés être des listes de chaînes
        # ne contiennent pas de dictionnaires, ce qui ferait planter l'interface React (Objects are not valid as a React child).
        for field in ["formations", "langues", "tools", "soft_skills", "technical_skills", "liens"]:
            if field in parsed_dict and isinstance(parsed_dict[field], list):
                new_list = []
                for item in parsed_dict[field]:
                    if isinstance(item, str):
                        new_list.append(item)
                    elif isinstance(item, dict):
                        new_list.append(" - ".join(str(v) for v in item.values() if v))
                    else:
                        new_list.append(str(item))
                parsed_dict[field] = new_list
                
        validated_data = CVExtraction(**parsed_dict)
        return json.dumps(validated_data.model_dump(), ensure_ascii=False, indent=4)
    except json.JSONDecodeError as e:
        print(f"[ERREUR] Le modèle Ollama n'a pas renvoyé un JSON valide: {raw_json_str}")
        raise ValueError(f"JSON invalide renvoyé par le modèle: {e}")
    except Exception as e:
        print(f"[ERREUR] Échec de la validation Pydantic : {e}")
        raise ValueError(f"Le JSON renvoyé ne correspond pas au schéma attendu : {e}")

